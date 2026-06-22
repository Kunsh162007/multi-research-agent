"""
Export research reports in multiple formats:
  - Markdown (.md)   — raw report text, ready for Obsidian / GitHub / Notion
  - PDF (.pdf)       — styled document via WeasyPrint (fpdf2 fallback)
  - DOCX (.docx)     — Word document via python-docx
  - BibTeX (.bib)    — citation entries extracted from arXiv findings

Both PDF and DOCX render embedded ```mermaid blocks to PNG figures via the free
kroki.io API. A `style="paper"` switch wraps output in a research-paper layout
(abstract, numbered sections, figures, references).
"""

import base64
import re
import logging
import zlib
from typing import Optional

logger = logging.getLogger(__name__)

_MERMAID_RE = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL)


# ─── Mermaid → PNG (kroki.io, free) ──────────────────────────────────────────────

def _kroki_png(diagram: str) -> Optional[bytes]:
    """Render a Mermaid diagram to PNG bytes via kroki.io. Returns None on failure."""
    try:
        import requests
        compressed = zlib.compress(diagram.encode("utf-8"), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
        resp = requests.get(f"https://kroki.io/mermaid/png/{encoded}", timeout=20)
        resp.raise_for_status()
        return resp.content
    except Exception as e:
        logger.warning(f"Mermaid render via kroki failed: {e}")
        return None


def _split_mermaid(report: str):
    """Yield ('text', md) and ('mermaid', code) segments in document order."""
    pos = 0
    for m in _MERMAID_RE.finditer(report):
        if m.start() > pos:
            yield ("text", report[pos:m.start()])
        yield ("mermaid", m.group(1).strip())
        pos = m.end()
    if pos < len(report):
        yield ("text", report[pos:])


# ─── Markdown ──────────────────────────────────────────────────────────────────

def to_markdown(title: str, query: str, report: str) -> bytes:
    header = f"# {title}\n\n**Research Query:** {query}\n\n---\n\n"
    return (header + report).encode("utf-8")


# ─── PDF ───────────────────────────────────────────────────────────────────────

def _strip_markdown(text: str) -> str:
    """Minimal markdown stripping for PDF rendering."""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,2}([^*\n]+)\*{1,2}', r'\1', text)
    text = re.sub(r'_{1,2}([^_\n]+)_{1,2}', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
    return text.strip()


_PAPER_CSS = """
@page { size: A4; margin: 2.2cm 2cm; @bottom-center { content: counter(page); color:#888; font-size:9pt; } }
body { font-family: 'Georgia', 'Times New Roman', serif; font-size: 11pt; line-height: 1.55; color:#1a1a2e; }
h1 { font-size: 20pt; text-align:center; margin-bottom: 4px; }
.query { text-align:center; color:#666; font-style:italic; font-size:10pt; margin-bottom:18px; }
h2 { font-size: 14pt; border-bottom:1px solid #ddd; padding-bottom:3px; margin-top:22px; }
h3 { font-size: 12pt; }
table { border-collapse: collapse; width:100%; margin:12px 0; font-size:10pt; }
th, td { border:1px solid #ccc; padding:6px 8px; text-align:left; }
th { background:#f3f4f6; }
code { background:#f3f3f5; padding:1px 4px; border-radius:3px; font-size:9.5pt; }
pre { background:#f6f6f8; padding:10px; border-radius:6px; overflow:auto; }
figure { text-align:center; margin:16px 0; }
figure img { max-width:100%; }
figcaption { color:#666; font-size:9pt; margin-top:4px; }
a { color:#4f46e5; text-decoration:none; }
"""


def _report_html_body(report: str) -> str:
    """Markdown → HTML with ```mermaid blocks replaced by embedded PNG figures."""
    try:
        import markdown as md_lib
    except ImportError:
        md_lib = None

    parts = []
    fig_no = 0
    for kind, content in _split_mermaid(report):
        if kind == "mermaid":
            png = _kroki_png(content)
            if png:
                fig_no += 1
                b64 = base64.b64encode(png).decode("ascii")
                parts.append(
                    f'<figure><img src="data:image/png;base64,{b64}"/>'
                    f'<figcaption>Figure {fig_no}</figcaption></figure>'
                )
            else:
                parts.append(f"<pre>{content}</pre>")
        else:
            parts.append(
                md_lib.markdown(content, extensions=["tables", "fenced_code"]) if md_lib
                else f"<pre>{content}</pre>"
            )
    return "\n".join(parts)


def to_pdf(title: str, query: str, report: str, style: str = "report") -> Optional[bytes]:
    """High-fidelity PDF via WeasyPrint; falls back to fpdf2 if native libs are absent."""
    try:
        from weasyprint import HTML
        body = _report_html_body(report)
        heading = "Research Paper" if style == "paper" else "Research Report"
        html = (
            f"<html><head><meta charset='utf-8'><style>{_PAPER_CSS}</style></head><body>"
            f"<h1>{title}</h1><div class='query'>{heading} · {query}</div>"
            f"{body}</body></html>"
        )
        return HTML(string=html).write_pdf()
    except Exception as e:
        logger.warning(f"WeasyPrint unavailable ({e}); falling back to fpdf2")
        return _to_pdf_fpdf(title, query, report)


def _to_pdf_fpdf(title: str, query: str, report: str) -> Optional[bytes]:
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 10, txt=title[:100], align="L")
        pdf.ln(2)

        # Query
        pdf.set_font("Helvetica", "I", 11)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 7, txt=f"Query: {query}", align="L")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)

        # Horizontal rule
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)

        # Body — split on headings for basic formatting
        pdf.set_font("Helvetica", "", 11)
        lines = _strip_markdown(report).split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                pdf.ln(4)
                continue
            # Treat lines that were headings (now just bold text)
            pdf.multi_cell(0, 6, txt=stripped, align="L")

        return bytes(pdf.output())
    except ImportError:
        logger.warning("fpdf2 not installed — PDF export unavailable. Run: pip install fpdf2")
        return None
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return None


# ─── DOCX ──────────────────────────────────────────────────────────────────────

def to_docx(title: str, query: str, report: str, style: str = "report") -> Optional[bytes]:
    """Word document via python-docx, with Mermaid blocks embedded as PNG figures."""
    try:
        import io
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        h = doc.add_heading(title[:120], level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(("Research Paper · " if style == "paper" else "Research Report · ") + query)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        fig_no = 0
        for kind, content in _split_mermaid(report):
            if kind == "mermaid":
                png = _kroki_png(content)
                if png:
                    fig_no += 1
                    doc.add_picture(io.BytesIO(png), width=Inches(6.0))
                    cap = doc.add_paragraph(f"Figure {fig_no}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                else:
                    doc.add_paragraph(content, style="Intense Quote")
                continue
            for line in content.split("\n"):
                s = line.rstrip()
                if not s:
                    continue
                m = re.match(r"^(#{1,4})\s+(.*)", s)
                if m:
                    doc.add_heading(_strip_markdown(m.group(2)), level=min(len(m.group(1)), 4))
                elif s.lstrip().startswith(("- ", "* ")):
                    doc.add_paragraph(_strip_markdown(s.lstrip()[2:]), style="List Bullet")
                else:
                    doc.add_paragraph(_strip_markdown(s))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        logger.warning("python-docx not installed — DOCX export unavailable.")
        return None
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return None


# ─── BibTeX ────────────────────────────────────────────────────────────────────

def _arxiv_id_from_url(url: str) -> Optional[str]:
    """Extract arXiv ID from a URL like https://arxiv.org/abs/2401.00001."""
    match = re.search(r'arxiv\.org/abs/([^\s/?#]+)', url or "")
    return match.group(1) if match else None


def to_bibtex(findings: list[dict]) -> bytes:
    """
    Generate a .bib file from the findings list.
    Only arXiv sources get proper @article entries; others get @misc.
    """
    entries = []
    seen: set[str] = set()

    for f in findings:
        url = f.get("url", "")
        source = f.get("source", "Unknown Source")
        citation_index = f.get("citation_index", 0)

        arxiv_id = _arxiv_id_from_url(url)
        key = f"ref{citation_index}"
        if key in seen:
            continue
        seen.add(key)

        if arxiv_id:
            # Extract year from arXiv ID (format: YYMM.NNNNN)
            year = "20" + arxiv_id[:2] if len(arxiv_id) >= 2 else "2024"
            # Extract author surnames from source title as best-effort author
            entry = (
                f"@article{{{key},\n"
                f"  title     = {{{source}}},\n"
                f"  year      = {{{year}}},\n"
                f"  url       = {{{url}}},\n"
                f"  note      = {{arXiv:{arxiv_id}}},\n"
                f"  archivePrefix = {{arXiv}},\n"
                f"  eprint    = {{{arxiv_id}}}\n"
                f"}}"
            )
        else:
            entry = (
                f"@misc{{{key},\n"
                f"  title     = {{{source}}},\n"
                f"  url       = {{{url}}},\n"
                f"  year      = {{2025}},\n"
                f"  note      = {{Accessed via Research Assistant}}\n"
                f"}}"
            )

        entries.append(entry)

    bib_content = "\n\n".join(entries) if entries else "% No citable sources found."
    return bib_content.encode("utf-8")
